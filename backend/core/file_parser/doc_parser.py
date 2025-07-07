"""Word文档解析器"""

from typing import List, Dict, Any, Optional, Callable, Generator
from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn
from loguru import logger

from prompt_hub import get_prompt

from .base_parser import BaseFileParser, ParsedContent
from core.model.vision.base import VisionModel
from .img_parser import ImgFileParser

# 全局命名空间字典，便于lxml查找
NS = {
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    # 可按需扩展
}

class WordFileParser(BaseFileParser):
    """Word文档解析器，支持文本和图片顺序还原"""
    
    def __init__(self, vision_model: VisionModel | None = None):
        """
        支持直接传入视觉模型实例（如OpenAIVisionModel等）
        """
        super().__init__()
        self.vision_model = vision_model
        if vision_model is not None:
            self.vision_model_func = vision_model.invoke
            self.img_parser = ImgFileParser(vision_model)
        else:
            self.vision_model_func = None
            self.img_parser = None
    
    def get_supported_extensions(self) -> List[str]:
        """获取支持的文件扩展名"""
        return ['docx']

    
    def parse_file_lazy(self, file_path: str | None = None, file_content: bytes | str | None = None) -> Generator[ParsedContent, None, None]:
        """
        懒加载方式解析Word文件内容，按文档实际顺序yield文本和图片
        """
        if not self.check_path(file_path):
            raise ValueError(f"不支持的文件格式或文件不存在: {file_path}")
        try:
            doc = Document(file_path)
            image_idx = 0
            rels = doc.part.rels
            image_rid_map = {}
            for rel in rels.values():
                if "image" in rel.target_ref:
                    image_rid_map[rel.rId] = rel.target_part.blob

            # 收集所有段落文本，便于图片上下文拼接
            all_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

            def yield_paragraph(paragraph: Paragraph, para_idx: int):
                nonlocal image_idx
                for run in paragraph.runs:
                    drawing_elements = run._element.findall('.//w:drawing', namespaces=NS)
                    if drawing_elements:
                        for drawing in drawing_elements:
                            blip = drawing.find('.//a:blip', namespaces=NS)
                            if blip is not None:
                                embed_rid = blip.attrib.get(qn('r:embed'))
                                if embed_rid and embed_rid in image_rid_map:
                                    image_data = image_rid_map[embed_rid]
                                    # 上下文：前后各1段
                                    context = ""
                                    if para_idx > 0:
                                        context += all_paragraphs[para_idx-1] + "\n"
                                    context += paragraph.text.strip() + "\n"
                                    if para_idx+1 < len(all_paragraphs):
                                        context += all_paragraphs[para_idx+1]
                                    metadata = {
                                        'image_idx': image_idx,
                                        'image_format': rels[embed_rid].target_ref.split('.')[-1] if '.' in rels[embed_rid].target_ref else 'unknown',
                                        'image_size': len(image_data)
                                    }
                                    if hasattr(self, 'img_parser') and self.img_parser is not None:
                                        # Use ImgParser to extract knowledge from image
                                        parsed = self.img_parser.parse_image(file_content=image_data, context=context)
                                        # Merge ImgParser metadata
                                        metadata.update(parsed.metadata)
                                        yield ParsedContent(
                                            content_type='image',
                                            content=f"【以下内容由视觉大模型解析的文档图片知识】\n{parsed.content}\n【文档图片知识结束】",
                                            metadata=metadata
                                        )
                                    else:
                                        # No vision model, yield placeholder
                                        yield ParsedContent(
                                            content_type='image',
                                            content='[Image Placeholder: This is an image in the document.]',
                                            metadata=metadata
                                        )
                                    image_idx += 1
                    # 普通文本run
                    text = run.text.strip()
                    if text:
                        style_name = paragraph.style.name if paragraph.style else ""
                        metadata = {
                            'style': style_name,
                            'alignment': str(paragraph.alignment) if paragraph.alignment else None,
                        }
                        if any(heading in style_name.lower() for heading in ['heading', 'title']):
                            content_type = 'heading'
                        else:
                            content_type = 'text'
                        yield ParsedContent(
                            content_type=content_type,
                            content=text,
                            metadata=metadata
                        )

            def yield_table(table: Table):
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            # 传递para_idx为-1，表格内不做上下文拼接
                            yield from yield_paragraph(para, -1)

            # 主体遍历
            for idx, element in enumerate(doc.element.body):
                if element.tag.endswith('p'):
                    para = self._get_paragraph_from_element(doc, element)
                    if para:
                        yield from yield_paragraph(para, idx)
                elif element.tag.endswith('tbl'):
                    table = self._get_table_from_element(doc, element)
                    if table:
                        yield from yield_table(table)
        except Exception as e:
            logger.error(f"Word文档流式解析失败: {file_path}, 错误: {str(e)}")
            raise ValueError(f"Word文档流式解析失败: {str(e)}")
    
    def _get_paragraph_from_element(self, doc: DocxDocument, element) -> Optional[Paragraph]:
        """从XML元素获取段落对象"""
        try:
            for para in doc.paragraphs:
                if para._element == element:
                    return para
            return None
        except Exception:
            return None
    
    def _get_table_from_element(self, doc: DocxDocument, element) -> Optional[Table]:
        """从XML元素获取表格对象"""
        try:
            for table in doc.tables:
                if table._element == element:
                    return table
            return None
        except Exception:
            return None
    
    def _parse_paragraph(self, paragraph: Paragraph) -> List[ParsedContent]:
        """
        解析段落内容
        
        Args:
            paragraph: 段落对象
            
        Returns:
            解析后的内容列表
        """
        contents = []
        text = paragraph.text.strip()
        
        if text:
            # 检查段落样式
            style_name = paragraph.style.name if paragraph.style else ""
            
            # 构建元数据
            metadata = {
                'style': style_name,
                'alignment': str(paragraph.alignment) if paragraph.alignment else None,
            }
            
            # 根据样式判断内容类型
            if any(heading in style_name.lower() for heading in ['heading', 'title']):
                content_type = 'heading'
            else:
                content_type = 'text'
            
            contents.append(ParsedContent(
                content_type=content_type,
                content=text,
                metadata=metadata
            ))
        
        return contents
    
    def _parse_table(self, table: Table) -> Optional[ParsedContent]:
        """
        解析表格内容
        
        Args:
            table: 表格对象
            
        Returns:
            解析后的表格内容
        """
        try:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                table_data.append(row_data)
            
            if not table_data:
                return None
            
            # 将表格转换为文本格式
            table_text = self._format_table_as_text(table_data)
            
            metadata = {
                'rows': len(table_data),
                'columns': len(table_data[0]) if table_data else 0,
                'table_data': table_data
            }
            
            return ParsedContent(
                content_type='table',
                content=table_text,
                metadata=metadata
            )
            
        except Exception as e:
            logger.warning(f"表格解析失败: {str(e)}")
            return None
    
    def _format_table_as_text(self, table_data: List[List[str]]) -> str:
        """将表格数据格式化为文本"""
        if not table_data:
            return ""
        
        # 简单的表格文本格式
        formatted_rows = []
        for row in table_data:
            formatted_rows.append(" | ".join(row))
        
        return "\n".join(formatted_rows)
    
